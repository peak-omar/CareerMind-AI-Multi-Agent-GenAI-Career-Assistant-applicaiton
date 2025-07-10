"""
CareerMind AI - Multi-Agent GenAI Career Assistant
Enhanced Data Loading and Processing Module with Advanced Features
"""

import os
import logging
import tempfile
from typing import Dict, List, Optional, Any, Union
from pathlib import Path
import hashlib
import json
from datetime import datetime

# Document processing imports
from docx import Document
from docx.shared import Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from langchain_community.document_loaders import PyMuPDFLoader
import PyPDF2
import io

# Data processing imports
import pandas as pd
import numpy as np
from PIL import Image
import base64

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


class EnhancedResumeLoader:
    """
    Enhanced resume loader with advanced parsing and analysis capabilities.
    """
    
    def __init__(self, cache_dir: str = "temp/cache"):
        """
        Initialize the enhanced resume loader.
        
        Args:
            cache_dir: Directory for caching processed files
        """
        self.cache_dir = Path(cache_dir)
        self.cache_dir.mkdir(parents=True, exist_ok=True)
        self.supported_formats = ['.pdf', '.docx', '.doc', '.txt']
        
    def load_resume(self, file_path: str) -> Dict[str, Any]:
        """
        Enhanced resume loading with comprehensive analysis.
        
        Args:
            file_path: Path to the resume file
            
        Returns:
            Dict containing resume content and metadata
        """
        try:
            file_path = Path(file_path)
            
            if not file_path.exists():
                raise FileNotFoundError(f"Resume file not found: {file_path}")
            
            # Check file format
            if file_path.suffix.lower() not in self.supported_formats:
                raise ValueError(f"Unsupported file format: {file_path.suffix}")
            
            # Generate cache key
            cache_key = self._generate_cache_key(file_path)
            cached_result = self._load_from_cache(cache_key)
            
            if cached_result:
                logger.info(f"Loaded resume from cache: {file_path.name}")
                return cached_result
            
            # Load and process the resume
            content = self._extract_content(file_path)
            metadata = self._extract_metadata(file_path)
            analysis = self._analyze_content(content)
            
            result = {
                'content': content,
                'metadata': metadata,
                'analysis': analysis,
                'file_info': {
                    'name': file_path.name,
                    'size': file_path.stat().st_size,
                    'format': file_path.suffix.lower(),
                    'processed_at': datetime.now().isoformat()
                }
            }
            
            # Cache the result
            self._save_to_cache(cache_key, result)
            
            logger.info(f"Successfully loaded and analyzed resume: {file_path.name}")
            return result
            
        except Exception as e:
            logger.error(f"Error loading resume {file_path}: {str(e)}")
            raise
    
    def _extract_content(self, file_path: Path) -> str:
        """
        Extract text content from various file formats.
        
        Args:
            file_path: Path to the file
            
        Returns:
            str: Extracted text content
        """
        try:
            if file_path.suffix.lower() == '.pdf':
                return self._extract_pdf_content(file_path)
            elif file_path.suffix.lower() in ['.docx', '.doc']:
                return self._extract_docx_content(file_path)
            elif file_path.suffix.lower() == '.txt':
                return self._extract_txt_content(file_path)
            else:
                raise ValueError(f"Unsupported file format: {file_path.suffix}")
                
        except Exception as e:
            logger.error(f"Error extracting content from {file_path}: {str(e)}")
            raise
    
    def _extract_pdf_content(self, file_path: Path) -> str:
        """
        Extract content from PDF files using multiple methods for better accuracy.
        
        Args:
            file_path: Path to PDF file
            
        Returns:
            str: Extracted text content
        """
        try:
            # Primary method: PyMuPDFLoader
            loader = PyMuPDFLoader(str(file_path))
            pages = loader.load()
            primary_content = "\n".join([page.page_content for page in pages])
            
            # Fallback method: PyPDF2
            if len(primary_content.strip()) < 100:  # If primary method didn't extract much
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    fallback_content = ""
                    for page in pdf_reader.pages:
                        fallback_content += page.extract_text() + "\n"
                    
                    if len(fallback_content.strip()) > len(primary_content.strip()):
                        return fallback_content
            
            return primary_content
            
        except Exception as e:
            logger.error(f"Error extracting PDF content: {str(e)}")
            raise
    
    def _extract_docx_content(self, file_path: Path) -> str:
        """
        Extract content from DOCX files.
        
        Args:
            file_path: Path to DOCX file
            
        Returns:
            str: Extracted text content
        """
        try:
            doc = Document(file_path)
            content = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    content.append(paragraph.text)
            
            # Extract table content
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            content.append(cell.text)
            
            return "\n".join(content)
            
        except Exception as e:
            logger.error(f"Error extracting DOCX content: {str(e)}")
            raise
    
    def _extract_txt_content(self, file_path: Path) -> str:
        """
        Extract content from text files.
        
        Args:
            file_path: Path to text file
            
        Returns:
            str: File content
        """
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
                
        except UnicodeDecodeError:
            # Try with different encoding
            with open(file_path, 'r', encoding='latin-1') as file:
                return file.read()
    
    def _extract_metadata(self, file_path: Path) -> Dict[str, Any]:
        """
        Extract metadata from the file.
        
        Args:
            file_path: Path to the file
            
        Returns:
            Dict containing file metadata
        """
        try:
            stat = file_path.stat()
            
            return {
                'file_size': stat.st_size,
                'created_at': datetime.fromtimestamp(stat.st_ctime).isoformat(),
                'modified_at': datetime.fromtimestamp(stat.st_mtime).isoformat(),
                'file_extension': file_path.suffix.lower(),
                'file_name': file_path.name
            }
            
        except Exception as e:
            logger.error(f"Error extracting metadata: {str(e)}")
            return {}
    
    def _analyze_content(self, content: str) -> Dict[str, Any]:
        """
        Perform basic content analysis on the resume.
        
        Args:
            content: Resume text content
            
        Returns:
            Dict containing content analysis results
        """
        try:
            words = content.split()
            lines = content.split('\n')
            
            # Basic statistics
            analysis = {
                'word_count': len(words),
                'line_count': len(lines),
                'character_count': len(content),
                'paragraph_count': len([line for line in lines if line.strip()]),
            }
            
            # Content detection
            content_lower = content.lower()
            
            # Detect sections
            sections = {
                'has_contact_info': any(keyword in content_lower for keyword in ['email', 'phone', 'linkedin', 'address']),
                'has_experience': any(keyword in content_lower for keyword in ['experience', 'work', 'employment', 'position']),
                'has_education': any(keyword in content_lower for keyword in ['education', 'university', 'college', 'degree']),
                'has_skills': any(keyword in content_lower for keyword in ['skills', 'technologies', 'programming', 'software']),
                'has_projects': any(keyword in content_lower for keyword in ['projects', 'portfolio', 'github']),
            }
            
            analysis['sections'] = sections
            analysis['completeness_score'] = sum(sections.values()) / len(sections) * 100
            
            return analysis
            
        except Exception as e:
            logger.error(f"Error analyzing content: {str(e)}")
            return {}
    
    def _generate_cache_key(self, file_path: Path) -> str:
        """
        Generate a cache key for the file.
        
        Args:
            file_path: Path to the file
            
        Returns:
            str: Cache key
        """
        try:
            stat = file_path.stat()
            content = f"{file_path.name}_{stat.st_size}_{stat.st_mtime}"
            return hashlib.md5(content.encode()).hexdigest()
            
        except Exception as e:
            logger.error(f"Error generating cache key: {str(e)}")
            return str(hash(str(file_path)))
    
    def _load_from_cache(self, cache_key: str) -> Optional[Dict[str, Any]]:
        """
        Load processed result from cache.
        
        Args:
            cache_key: Cache key
            
        Returns:
            Dict if cached result exists, None otherwise
        """
        try:
            cache_file = self.cache_dir / f"{cache_key}.json"
            
            if cache_file.exists():
                with open(cache_file, 'r', encoding='utf-8') as f:
                    return json.load(f)
            
            return None
            
        except Exception as e:
            logger.error(f"Error loading from cache: {str(e)}")
            return None
    
    def _save_to_cache(self, cache_key: str, result: Dict[str, Any]) -> None:
        """
        Save processed result to cache.
        
        Args:
            cache_key: Cache key
            result: Result to cache
        """
        try:
            cache_file = self.cache_dir / f"{cache_key}.json"
            
            with open(cache_file, 'w', encoding='utf-8') as f:
                json.dump(result, f, indent=2, ensure_ascii=False)
            
        except Exception as e:
            logger.error(f"Error saving to cache: {str(e)}")


class EnhancedCoverLetterGenerator:
    """
    Enhanced cover letter generator with multiple templates and customization options.
    """
    
    def __init__(self, template_dir: str = "temp/templates"):
        """
        Initialize the cover letter generator.
        
        Args:
            template_dir: Directory for storing templates
        """
        self.template_dir = Path(template_dir)
        self.template_dir.mkdir(parents=True, exist_ok=True)
        self.templates = self._load_templates()
    
    def generate_cover_letter(
        self,
        content: str,
        company_name: str = "Company",
        position: str = "Position",
        template_style: str = "professional",
        custom_formatting: Optional[Dict[str, Any]] = None
    ) -> str:
        """
        Generate an enhanced cover letter with custom formatting.
        
        Args:
            content: Cover letter content
            company_name: Target company name
            position: Job position
            template_style: Style template to use
            custom_formatting: Custom formatting options
            
        Returns:
            str: Path to generated cover letter file
        """
        try:
            # Create new document
            doc = Document()
            
            # Apply custom styling
            self._apply_document_styling(doc, template_style, custom_formatting)
            
            # Add header
            self._add_header(doc, company_name, position)
            
            # Process content
            paragraphs = content.split('\n')
            
            for paragraph in paragraphs:
                if paragraph.strip():
                    p = doc.add_paragraph(paragraph.strip())
                    self._style_paragraph(p, template_style)
            
            # Add footer
            self._add_footer(doc)
            
            # Generate filename
            safe_company = "".join(c for c in company_name if c.isalnum() or c in (' ', '-', '_')).rstrip()
            filename = f"temp/{safe_company}_{position}_cover_letter_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
            
            # Save document
            doc.save(filename)
            
            logger.info(f"Enhanced cover letter generated: {filename}")
            return filename
            
        except Exception as e:
            logger.error(f"Error generating cover letter: {str(e)}")
            raise
    
    def _apply_document_styling(
        self, 
        doc: Document, 
        template_style: str, 
        custom_formatting: Optional[Dict[str, Any]]
    ) -> None:
        """
        Apply document-level styling.
        
        Args:
            doc: Document object
            template_style: Style template name
            custom_formatting: Custom formatting options
        """
        try:
            # Set document margins
            sections = doc.sections
            for section in sections:
                section.top_margin = Inches(1)
                section.bottom_margin = Inches(1)
                section.left_margin = Inches(1)
                section.right_margin = Inches(1)
            
            # Apply template-specific styling
            if template_style == "professional":
                # Professional styling
                pass
            elif template_style == "modern":
                # Modern styling with colors
                pass
            elif template_style == "creative":
                # Creative styling
                pass
            
            # Apply custom formatting if provided
            if custom_formatting:
                # Apply custom styles
                pass
                
        except Exception as e:
            logger.error(f"Error applying document styling: {str(e)}")
    
    def _add_header(self, doc: Document, company_name: str, position: str) -> None:
        """
        Add a professional header to the document.
        
        Args:
            doc: Document object
            company_name: Company name
            position: Position title
        """
        try:
            # Add date
            date_paragraph = doc.add_paragraph()
            date_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            date_run = date_paragraph.add_run(datetime.now().strftime("%B %d, %Y"))
            
            # Add spacing
            doc.add_paragraph()
            
            # Add recipient info
            if company_name != "Company":
                doc.add_paragraph(f"Hiring Manager\n{company_name}")
                doc.add_paragraph()
            
            # Add subject line
            if position != "Position":
                subject_para = doc.add_paragraph()
                subject_run = subject_para.add_run(f"Subject: Application for {position}")
                subject_run.bold = True
                doc.add_paragraph()
            
        except Exception as e:
            logger.error(f"Error adding header: {str(e)}")
    
    def _style_paragraph(self, paragraph, template_style: str) -> None:
        """
        Apply styling to a paragraph based on template.
        
        Args:
            paragraph: Paragraph object
            template_style: Style template name
        """
        try:
            if template_style == "professional":
                paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            elif template_style == "modern":
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
        except Exception as e:
            logger.error(f"Error styling paragraph: {str(e)}")
    
    def _add_footer(self, doc: Document) -> None:
        """
        Add a professional footer to the document.
        
        Args:
            doc: Document object
        """
        try:
            # Add closing
            doc.add_paragraph()
            doc.add_paragraph("Sincerely,")
            doc.add_paragraph()
            doc.add_paragraph("[Your Name]")
            
        except Exception as e:
            logger.error(f"Error adding footer: {str(e)}")
    
    def _load_templates(self) -> Dict[str, Dict[str, Any]]:
        """
        Load available cover letter templates.
        
        Returns:
            Dict containing template configurations
        """
        return {
            "professional": {
                "font": "Times New Roman",
                "font_size": 11,
                "line_spacing": 1.15,
                "paragraph_spacing": 6
            },
            "modern": {
                "font": "Calibri",
                "font_size": 11,
                "line_spacing": 1.2,
                "paragraph_spacing": 8
            },
            "creative": {
                "font": "Arial",
                "font_size": 10,
                "line_spacing": 1.1,
                "paragraph_spacing": 5
            }
        }


# Backward compatibility functions
def load_resume(file_path: str) -> str:
    """
    Backward compatible resume loading function.
    
    Args:
        file_path: Path to resume file
        
    Returns:
        str: Resume content
    """
    try:
        loader = EnhancedResumeLoader()
        result = loader.load_resume(file_path)
        return result['content']
        
    except Exception as e:
        logger.error(f"Error in backward compatible load_resume: {str(e)}")
        # Fallback to simple loading
        if file_path.endswith('.pdf'):
            loader = PyMuPDFLoader(file_path)
            pages = loader.load()
            return "\n".join([page.page_content for page in pages])
        else:
            raise


def write_cover_letter_to_doc(
    text: str, 
    filename: str = "temp/cover_letter.docx",
    company_name: str = "Company"
) -> str:
    """
    Enhanced backward compatible cover letter writing function.
    
    Args:
        text: Cover letter text
        filename: Output filename
        company_name: Company name for enhanced formatting
        
    Returns:
        str: Path to generated file
    """
    try:
        generator = EnhancedCoverLetterGenerator()
        return generator.generate_cover_letter(
            content=text,
            company_name=company_name,
            template_style="professional"
        )
        
    except Exception as e:
        logger.error(f"Error in enhanced cover letter generation: {str(e)}")
        # Fallback to simple generation
        doc = Document()
        paragraphs = text.split("\n")
        for para in paragraphs:
            if para.strip():
                doc.add_paragraph(para.strip())
        doc.save(filename)
        return filename


# Utility functions
def validate_file_format(file_path: str) -> bool:
    """
    Validate if the file format is supported.
    
    Args:
        file_path: Path to file
        
    Returns:
        bool: True if format is supported
    """
    supported_formats = ['.pdf', '.docx', '.doc', '.txt']
    return Path(file_path).suffix.lower() in supported_formats


def get_file_info(file_path: str) -> Dict[str, Any]:
    """
    Get comprehensive file information.
    
    Args:
        file_path: Path to file
        
    Returns:
        Dict containing file information
    """
    try:
        path = Path(file_path)
        stat = path.stat()
        
        return {
            'name': path.name,
            'size': stat.st_size,
            'size_mb': round(stat.st_size / (1024 * 1024), 2),
            'extension': path.suffix.lower(),
            'created': datetime.fromtimestamp(stat.st_ctime).isoformat(),
            'modified': datetime.fromtimestamp(stat.st_mtime).isoformat(),
            'is_supported': validate_file_format(file_path)
        }
        
    except Exception as e:
        logger.error(f"Error getting file info: {str(e)}")
        return {}